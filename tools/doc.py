from collections.abc import Generator
import os
import tempfile
import io
import markdown
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import Any, List, Dict, Tuple
import html
from bs4 import BeautifulSoup, NavigableString, Tag

from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage

class DocTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        md_content = tool_parameters.get("markdown_content", "")
        title = tool_parameters.get("title", "Document")
        
        if not md_content:
            yield self.create_text_message("No markdown content provided.")
            return
        
        try:
            # 只去除“↓”，保留“•”
            md_content = md_content.replace("↓", "")
            
            doc = self._convert_markdown_to_docx(md_content, title)
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            file_bytes = docx_bytes.getvalue()
            filename = f"{title.replace(' ', '_')}.docx"
            
            yield self.create_text_message(f"Document '{title}' generated successfully")
            yield self.create_blob_message(
                blob=file_bytes,
                meta={
                    "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "filename": filename
                }
            )
        except Exception as e:
            yield self.create_text_message(f"Error converting markdown to DOCX: {str(e)}")

    def _convert_markdown_to_docx(self, md_content: str, title: str) -> Document:
        doc = Document()
        
        # --- 标题部分 ---
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.italic = False
        title_run.font.name = '黑体'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        title_run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
        title_run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        title_run.font.size = Pt(22)
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 转换 Markdown 为 HTML（移除 nl2br 扩展）
        html_content = markdown.markdown(
            md_content,
            extensions=[
                'markdown.extensions.tables',
                'markdown.extensions.fenced_code',
                'markdown.extensions.codehilite',
                # 'markdown.extensions.nl2br',
                'markdown.extensions.sane_lists'
            ]
        )
        
        soup = BeautifulSoup(html_content, 'html.parser')
        self._process_html_elements(doc, soup)
        return doc
    
    def _is_chart_related_text(self, text: str) -> bool:
        """
        检查文本是否包含图表相关关键词
        """
        if not text or not text.strip():
            return False
        
        text_lower = text.lower().strip()
        chart_keywords = ['echarts', 'chart', 'graph', 'plot', 'canvas', 'visualization', 
                         'highcharts', 'd3', 'chartjs', 'amcharts', 'plotly']
        
        return any(keyword in text_lower for keyword in chart_keywords)

    def _process_html_elements(self, doc: Document, soup: BeautifulSoup) -> None:
        """
        递归遍历所有 HTML 节点，根据标签名称将内容插入到 Word 文档中。
        只去除“↓”，保留“•”。
        """
        for element in soup.children:
            # 跳过 <br> 标签
            if isinstance(element, Tag) and element.name == 'br':
                continue
            
            # 纯文本节点：只去"↓"
            if isinstance(element, NavigableString):
                text = str(element).replace("↓", "").strip()
                # 过滤图表相关文本
                if self._is_chart_related_text(text):
                    continue
                # 跳过短文本（如"分析："），避免它们成为独立段落
                # 这些短文本通常是标签，应该与后续内容合并
                if text and len(text) > 10:  # 只处理较长的文本
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Pt(24)  # 2个字符缩进，与正文段落一致
                    p.paragraph_format.line_spacing = 1.25  # 1.25倍行间距
                    run = p.add_run(text)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    run.font.size = Pt(12)  # 小四
                    run.font.color.rgb = RGBColor(0, 0, 0)
                continue
            
            if not isinstance(element, Tag):
                continue
            
            # 处理标题 h1 ~ h6，只去除“↓”
            if element.name == 'h1':
                heading = doc.add_heading(element.get_text().replace("↓", "").strip(), level=1)
                for run in heading.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    run.font.size = Pt(16)  # 三号
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.italic = False

            elif element.name == 'h2':
                heading = doc.add_heading(element.get_text().replace("↓", "").strip(), level=2)
                for run in heading.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    run.font.size = Pt(14)  # 四号
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.italic = False

            elif element.name == 'h3':
                heading = doc.add_heading(element.get_text().replace("↓", "").strip(), level=3)
                for run in heading.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    run.font.size = Pt(12)  # 小四
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.italic = False

            elif element.name == 'h4':
                heading = doc.add_heading(element.get_text().replace("↓", "").strip(), level=4)
                for run in heading.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    run.font.size = Pt(12)  # 小四
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.italic = False

            elif element.name == 'h5':
                heading = doc.add_heading(element.get_text().replace("↓", "").strip(), level=5)
                for run in heading.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    run.font.size = Pt(12)  # 小四
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.italic = False

            elif element.name == 'h6':
                heading = doc.add_heading(element.get_text().replace("↓", "").strip(), level=6)
                for run in heading.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    run.font.size = Pt(12)  # 小四
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.italic = False

            # 段落 <p>：只去“↓”
            elif element.name == 'p':
                text = element.get_text().replace("↓", "").strip()
                if not text:
                    continue

                if text.startswith('附件'):
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Pt(24)  # 2个字符缩进
                    p.paragraph_format.line_spacing = 1.25  # 1.25倍行间距
                    self._add_run_with_formatting(p, element)
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Pt(24)  # 2个字符缩进
                    p.paragraph_format.line_spacing = 1.25  # 1.25倍行间距
                    self._add_run_with_formatting(p, element)

            # 列表 <ul> / <ol>：只去“↓”，保留“•”
            elif element.name == 'ul':
                self._add_list(doc, element, is_numbered=False)

            elif element.name == 'ol':
                self._add_list(doc, element, is_numbered=True)

            # 代码块 <pre>：只去“↓”
            elif element.name == 'pre':
                code = element.get_text().replace("↓", "")
                lang = ""
                if element.code and element.code.get('class'):
                    for cls in element.code.get('class'):
                        if cls.startswith('language-'):
                            lang = cls[9:]
                            break
                self._add_code_block(doc, code, lang)

            # 表格 <table>：只去"↓"
            elif element.name == 'table':
                self._add_html_table(doc, element)

            # 分隔线 <hr>
            elif element.name == 'hr':
                doc.add_paragraph('_' * 50)

            # 忽略图表和脚本相关标签
            elif element.name in ['script', 'style', 'canvas']:
                continue
            
            # 忽略包含图表ID或类名的div标签
            elif element.name == 'div':
                div_id = element.get('id', '').lower()
                div_class = ' '.join(element.get('class', [])).lower()
                div_text = element.get_text().strip().lower()
                
                # 如果div包含图表相关关键词，则忽略
                chart_keywords = [
                    'echarts', 'echart', 'chart', 'charts', 'graph', 'plot', 'canvas',
                    'visualization', 'visualize', 'highcharts', 'highchart',
                    'd3', 'plotly', 'chartjs', 'amcharts', '图表', '图表展示'
                ]
                if any(keyword in div_id or keyword in div_class or keyword in div_text for keyword in chart_keywords):
                    continue
                else:
                    # 普通div，递归处理其子元素
                    self._process_html_elements(doc, element)

            # 其他容器标签（例如 <span> 等），递归处理其子元素
            else:
                self._process_html_elements(doc, element)

    def _add_run_with_formatting(self, paragraph, element):
        """
        在一个段落里插入富文本内容，并保持“只去除↓”的原则。
        """
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child).replace("↓", "")
                # 过滤图表相关文本
                if self._is_chart_related_text(text):
                    continue
                run = paragraph.add_run(text)

            elif isinstance(child, Tag) and child.name in ['strong', 'b']:
                text = child.get_text().replace("↓", "")
                # 过滤图表相关文本
                if self._is_chart_related_text(text):
                    continue
                run = paragraph.add_run(text)

            elif isinstance(child, Tag) and child.name in ['em', 'i']:
                text = child.get_text().replace("↓", "")
                # 过滤图表相关文本
                if self._is_chart_related_text(text):
                    continue
                run = paragraph.add_run(text)
                run.italic = True

            elif isinstance(child, Tag) and child.name == 'code':
                code_text = child.get_text().replace("↓", "")
                run = paragraph.add_run(code_text)
                run.font.name = 'Courier New'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
                run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 0, 0)

            elif isinstance(child, Tag):
                # 递归处理其他标签（<span>、<a>、<div> 等）
                self._add_run_with_formatting(paragraph, child)
            
            # 统一为正文文字格式：宋体，12pt，黑色
            try:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                run.font.size = Pt(12)  # 小四
                run.font.color.rgb = RGBColor(0, 0, 0)
            except UnboundLocalError:
                # 如果 run 未定义（极端情况），忽略即可
                pass

    def _add_list(self, doc: Document, list_element: Tag, is_numbered: bool = False):
        """
        处理 HTML 中的 <ul> 或 <ol> 列表，只去“↓”，保留项目符号。
        """
        for item in list_element.find_all('li', recursive=False):
            # 提取 li 文本并只去"↓"
            item_text = "".join(t for t in item.strings).replace("↓", "").strip()
            if not item_text:
                continue
            
            # 过滤图表相关文本
            if self._is_chart_related_text(item_text):
                continue

            # 根据是否有序，选择不同的样式
            p = doc.add_paragraph(style='List Number' if is_numbered else 'List Bullet')
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.25  # 1.25倍行间距
            run = p.add_run(item_text)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
            run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            run.font.size = Pt(12)  # 小四
            run.font.color.rgb = RGBColor(0, 0, 0)

            # 递归处理嵌套的 <ul> / <ol>
            nested_ul = item.find('ul')
            nested_ol = item.find('ol')
            if nested_ul:
                self._add_list(doc, nested_ul, is_numbered=False)
            if nested_ol:
                self._add_list(doc, nested_ol, is_numbered=True)

    def _add_code_block(self, doc: Document, code: str, language: str = ""):
        """
        将 <pre><code>…</code></pre> 转为 Word 中的代码块格式，只去“↓”。
        """
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing = 1.25  # 1.25倍行间距
        if language:
            lang_run = p.add_run(f"Language: {language}\n")
            lang_run.italic = True
            lang_run.font.name = '宋体'
            lang_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            lang_run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
            lang_run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            lang_run.font.size = Pt(12)  # 小四
            lang_run.font.color.rgb = RGBColor(0, 0, 0)

        code_filtered = code.replace("↓", "")
        code_run = p.add_run(code_filtered)
        code_run.font.name = 'Courier New'
        code_run.font.size = Pt(9)
        code_run.font.color.rgb = RGBColor(0, 0, 0)

        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)

    def _add_html_table(self, doc: Document, table: Tag):
        """
        将 HTML 中的 <table> 转为 Word 表格，只去“↓”。
        """
        rows = table.find_all('tr')
        if not rows:
            return

        # 第一行决定列数
        header_cells = rows[0].find_all(['th', 'td'])
        col_count = len(header_cells)
        doc_table = doc.add_table(rows=len(rows), cols=col_count)
        doc_table.style = 'Table Grid'

        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            for j, cell in enumerate(cells):
                if j >= col_count:
                    continue
                # 只去"↓"，保留其他符号
                cell_text = cell.get_text().replace("↓", "")
                
                # 过滤图表相关文本
                if self._is_chart_related_text(cell_text):
                    cell_text = ""  # 清空图表相关内容
                
                doc_table.cell(i, j).text = cell_text

                for para in doc_table.cell(i, j).paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
                    for run in para.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                        run.font.size = Pt(12)  # 小四
                        run.font.color.rgb = RGBColor(0, 0, 0)

                # 如果是表头行（i == 0）或单元格标签是 <th>，设置居中对齐
                if i == 0 or cell.name == 'th':
                    for para in doc_table.cell(i, j).paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 表头也居中对齐
                        for run in para.runs:
                            run.font.color.rgb = RGBColor(0, 0, 0)
    
    def _is_chart_related_text(self, text: str) -> bool:
        """
        检查文本是否与图表相关，需要过滤掉
        """
        if not text:
            return False
        
        text_lower = text.lower().strip()
        chart_keywords = [
            'echarts', 'echart', 'chart', 'charts', '图表', '图表说明', '图表描述', '图表展示',
            'canvas', 'svg', 'visualization', '可视化', 'visualize',
            '图例', '坐标轴', 'tooltip', '数据可视化', 'highcharts', 'highchart',
            'd3.js', 'd3', 'plotly', 'chartjs', 'chart.js', 'amcharts',
            '柱状图', '折线图', '饼图', '散点图', '条形图', '雷达图', '图表展示：'
        ]
        
        return any(keyword in text_lower for keyword in chart_keywords)
    
    def _reset_number_counter(self):
        """
        重置编号计数器，在每个标题后调用
        """
        self.number_counter = 1
        self.number_mapping = {}
    
    def _convert_number_labels(self, text: str) -> str:
        """
        将标题下的编号格式转换为带括号的格式，以区别于标题编号
        例如：将 "1. 内容" 转换为 "(1) 内容"
        """
        import re
        
        def replace_number(match):
            original_num = match.group(1)
            if original_num not in self.number_mapping:
                self.number_mapping[original_num] = str(self.number_counter)
                self.number_counter += 1
            new_num = self.number_mapping[original_num]
            return f"({new_num}) "
        
        # 匹配各种编号格式：1. 、1、、1）、（1）
        patterns = [
            r'(\d+)\.\s*',      # 1. 
            r'(\d+)、\s*',      # 1、
            r'(\d+)\)\s*',      # 1)
            r'\((\d+)\)\s*'     # (1)
        ]
        
        result = text
        for pattern in patterns:
            result = re.sub(pattern, replace_number, result)
        
        return result
